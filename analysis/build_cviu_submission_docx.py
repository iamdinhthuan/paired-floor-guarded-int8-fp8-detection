#!/usr/bin/env python3
"""Build the two editable CVIU submission documents from plain-text sources."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt


AUTHOR = "Dinh Thuan Nguyen"


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    document.core_properties.author = AUTHOR
    document.core_properties.last_modified_by = AUTHOR
    document.core_properties.created = datetime(2026, 9, 3, 8, 0, 0)
    document.core_properties.modified = datetime(2026, 9, 3, 8, 0, 0)
    document.core_properties.comments = (
        "Generated deterministically from the checked plain-text submission source."
    )


def build_highlights(source: Path, output: Path) -> None:
    items = [line[2:].strip() for line in source.read_text(encoding="utf-8").splitlines() if line.startswith("- ")]
    if not 3 <= len(items) <= 5:
        raise ValueError(f"Elsevier Highlights require 3--5 items, found {len(items)}")
    too_long = [(item, len(item)) for item in items if len(item) > 85]
    if too_long:
        raise ValueError(f"Highlights exceed 85 characters: {too_long}")

    document = Document()
    configure(document)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Highlights")
    run.bold = True
    run.font.size = Pt(14)
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(item)
    document.save(output)


def build_cover_letter(source: Path, output: Path) -> None:
    blocks = [block.strip() for block in source.read_text(encoding="utf-8").split("\n\n") if block.strip()]
    document = Document()
    configure(document)
    for index, block in enumerate(blocks):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for line_index, line in enumerate(block.splitlines()):
            if line_index:
                paragraph.add_run().add_break()
            paragraph.add_run(line)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--paper-root", type=Path, default=Path("paper"))
    args = parser.parse_args()
    root = args.paper_root.resolve()
    build_highlights(root / "highlights.txt", root / "Highlights.docx")
    build_cover_letter(root / "cover_letter.txt", root / "Cover_Letter.docx")
    print(f"wrote {root / 'Highlights.docx'}")
    print(f"wrote {root / 'Cover_Letter.docx'}")


if __name__ == "__main__":
    main()
